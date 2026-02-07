"""
Script để đọc báo cáo gốc và tạo file DOCX mới với sơ đồ khối
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Đọc file gốc để hiểu format
try:
    doc_original = Document('Bao_cao_nghien_cuu_de_tai.docx')
    print("=== NỘI DUNG BÁO CÁO GỐC ===")
    for i, para in enumerate(doc_original.paragraphs[:20]):
        if para.text.strip():
            print(f"[{i}] Style: {para.style.name} | Text: {para.text[:100]}...")
except Exception as e:
    print(f"Lỗi đọc file: {e}")

# Tạo document mới
doc = Document()

# Set font mặc định
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)

# ========== TIÊU ĐỀ CHƯƠNG ==========
title = doc.add_heading('', level=1)
run = title.add_run('CHƯƠNG 3: PHÂN TÍCH THIẾT KẾ HỆ THỐNG')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ========== PHẦN 3.1 ==========
doc.add_heading('3.1. Sơ đồ khối quy trình huấn luyện mô hình', level=2)

doc.add_paragraph(
    'Hình dưới đây mô tả chi tiết quy trình xây dựng và huấn luyện mô hình nhận diện cảm xúc. '
    'Dữ liệu ảnh được chia thành 3 tập: Train, Validation và Test. Sau đó đi qua các bước tiền xử lý, '
    'tăng cường dữ liệu, huấn luyện với kiến trúc MobileNetV2, và cuối cùng đánh giá kết quả.'
)

# ========== VẼ SƠ ĐỒ BẰNG TABLE (Thay thế cho hình ảnh) ==========
doc.add_paragraph('')  # Khoảng trắng

# Tạo bảng làm sơ đồ
table = doc.add_table(rows=9, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Hàm helper để set cell
def set_cell(cell, text, bold=False, center=True):
    cell.text = text
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold

# Row 0: Dataset
set_cell(table.cell(0, 1), '📁 BỘ DỮ LIỆU ẢNH KHUÔN MẶT\n(7 lớp cảm xúc)', bold=True)
table.cell(0, 0).merge(table.cell(0, 0))
table.cell(0, 2).merge(table.cell(0, 2))

# Row 1: Arrow
set_cell(table.cell(1, 1), '↓')

# Row 2: Split data
set_cell(table.cell(2, 0), 'Tập TRAIN\n(80%)', bold=True)
set_cell(table.cell(2, 1), 'Tập VALIDATION\n(20%)', bold=True)
set_cell(table.cell(2, 2), 'Tập TEST\n(Độc lập)', bold=True)

# Row 3: Arrow
set_cell(table.cell(3, 0), '↓')
set_cell(table.cell(3, 1), '↓')
set_cell(table.cell(3, 2), '↓')

# Row 4: Preprocessing
set_cell(table.cell(4, 0), '🔄 TIỀN XỬ LÝ\n• Resize 48x48 RGB\n• Chuẩn hóa [0,1]\n• Data Augmentation', bold=True)
set_cell(table.cell(4, 1), '🔄 TIỀN XỬ LÝ\n• Resize 48x48 RGB\n• Chuẩn hóa [0,1]', bold=True)
set_cell(table.cell(4, 2), '🔄 TIỀN XỬ LÝ\n• Resize 48x48 RGB\n• Chuẩn hóa [0,1]', bold=True)

# Row 5: Arrow
set_cell(table.cell(5, 0), '↓')
set_cell(table.cell(5, 1), '↓')
set_cell(table.cell(5, 2), '')

# Row 6: Model
table.cell(6, 0).merge(table.cell(6, 1))
set_cell(table.cell(6, 0), '🧠 MÔ HÌNH MobileNetV2\n• Transfer Learning (ImageNet)\n• Fine-tune 30 lớp cuối\n• Dense 256 → 128 → 7 (Softmax)', bold=True)
set_cell(table.cell(6, 2), '(Chờ đánh giá)')

# Row 7: Arrow
set_cell(table.cell(7, 0), '↓')
set_cell(table.cell(7, 1), '')
set_cell(table.cell(7, 2), '↓')

# Row 8: Output
table.cell(8, 0).merge(table.cell(8, 1))
set_cell(table.cell(8, 0), '✅ MÔ HÌNH TỐI ƯU\n(best_model.keras)', bold=True)
set_cell(table.cell(8, 2), '📊 ĐÁNH GIÁ\nAccuracy, F1-Score\nConfusion Matrix', bold=True)

# Caption
caption = doc.add_paragraph('Hình 3.1. Sơ đồ khối quy trình huấn luyện mô hình')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.runs[0].italic = True

# ========== PHẦN 3.2 ==========
doc.add_heading('3.2. Mô tả chi tiết các khối chức năng', level=2)

# Khối 1
doc.add_heading('3.2.1. Khối Dữ liệu (Data Block)', level=3)
doc.add_paragraph(
    'Bộ dữ liệu gồm các ảnh khuôn mặt đã được gán nhãn thuộc 7 lớp cảm xúc: '
    'Angry (Giận dữ), Disgust (Ghê tởm), Fear (Sợ hãi), Happy (Hạnh phúc), '
    'Neutral (Bình thường), Sad (Buồn), Surprise (Ngạc nhiên). '
    'Dữ liệu được chia theo tỷ lệ 80% Train - 20% Validation, và một tập Test riêng biệt.'
)

# Khối 2
doc.add_heading('3.2.2. Khối Tiền xử lý và Tăng cường dữ liệu', level=3)
p = doc.add_paragraph('Các bước tiền xử lý bao gồm:')
doc.add_paragraph('Resize ảnh về kích thước 48x48 pixel (3 kênh màu RGB).', style='List Bullet')
doc.add_paragraph('Chuẩn hóa giá trị pixel về khoảng [0, 1] bằng cách chia cho 255.', style='List Bullet')
doc.add_paragraph('Data Augmentation (chỉ áp dụng cho tập Train): Xoay ngẫu nhiên, dịch chuyển ngang/dọc, lật ngang, thay đổi độ sáng.', style='List Bullet')

# Khối 3
doc.add_heading('3.2.3. Khối Mô hình (Model Architecture)', level=3)
doc.add_paragraph(
    'Sử dụng kiến trúc MobileNetV2 với kỹ thuật Transfer Learning. Base model được load trọng số từ ImageNet, '
    'đóng băng các lớp đầu và fine-tune 30 lớp cuối. Phần Classification Head gồm: '
    'GlobalAveragePooling2D → Dense(256) → BatchNorm → Dropout(0.5) → Dense(128) → Dense(7, Softmax).'
)

# Khối 4
doc.add_heading('3.2.4. Khối Huấn luyện và Đánh giá', level=3)
doc.add_paragraph(
    'Quá trình huấn luyện sử dụng optimizer Adam (learning rate = 0.0001), hàm loss Categorical Crossentropy '
    'với Label Smoothing. Các Callbacks được sử dụng: EarlyStopping (patience=10), ReduceLROnPlateau, '
    'ModelCheckpoint để lưu model tốt nhất dựa trên val_accuracy. Sau khi train, model được đánh giá trên tập Test '
    'bằng các chỉ số Accuracy, Precision, Recall, F1-Score và Confusion Matrix.'
)

# ========== LƯU FILE ==========
output_path = 'Chuong3_Phan_tich_thiet_ke.docx'
doc.save(output_path)
print(f"\n✅ Đã tạo file: {output_path}")
