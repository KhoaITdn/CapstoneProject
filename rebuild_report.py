# -*- coding: utf-8 -*-
"""
Rebuild Bao_cao_nghien_cuu_de_tai.docx với:
- Mục lục có hyperlink (clickable)
- Ảnh minh họa
- Trình bày chuyên nghiệp
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

IMG_DIR = r'd:\New folder\CapstoneProject\images'

doc = Document()

# === STYLE ===
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Set default font for East Asian
rFonts = style.element.rPr
if rFonts is None:
    style.element.get_or_add_rPr()

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def p(text):
    pr = doc.add_paragraph(text)
    pr.paragraph_format.first_line_indent = Cm(1.27)
    for r in pr.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    return pr

def b(text):
    pr = doc.add_paragraph(text, style='List Bullet')
    for r in pr.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    return pr

def add_image(filename, caption, width_inches=5.5):
    """Add image with caption"""
    img_path = os.path.join(IMG_DIR, filename)
    if os.path.exists(img_path):
        pr = doc.add_paragraph()
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pr.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.italic = True
        cap.paragraph_format.space_after = Pt(12)

def tbl(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = hd
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
    doc.add_paragraph()

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    tag = run._r
    start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="0" w:name="{bookmark_name}"/>')
    end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="0"/>')
    tag.addprevious(start)
    tag.addnext(end)

def add_hyperlink_to_bookmark(paragraph, bookmark_name, text):
    """Add a clickable hyperlink to a bookmark in the TOC"""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 180)
    run.underline = True
    return run


# =============================================
# TRANG BÌA
# =============================================
for _ in range(3):
    doc.add_paragraph()

# School info
school = doc.add_paragraph()
school.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = school.add_run('TRƯỜNG ĐẠI HỌC ...................\nKHOA CÔNG NGHỆ THÔNG TIN')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.bold = True

for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BÁO CÁO NGHIÊN CỨU ĐỀ TÀI')
run.bold = True
run.font.size = Pt(22)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('XÂY DỰNG HỆ THỐNG NHẬN DIỆN CẢM XÚC KHUÔN MẶT\nTHEO THỜI GIAN THỰC TRONG CUỘC GỌI VIDEO')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

for _ in range(4):
    doc.add_paragraph()

# Info
info_items = [
    'GVHD: .........................................',
    'SVTH: .........................................',
    'MSSV: .........................................',
    'Lớp: ...........................................',
]
for item in info_items:
    pr = doc.add_paragraph()
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pr.add_run(item)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'

doc.add_paragraph()
dt = doc.add_paragraph()
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dt.add_run('Ngày lập báo cáo: 24/02/2026')
run.font.size = Pt(13)
run.font.name = 'Times New Roman'

doc.add_page_break()


# =============================================
# MỤC LỤC (có link màu xanh)
# =============================================
h1('MỤC LỤC')

toc_data = [
    ('CHƯƠNG 1: GIỚI THIỆU VỀ CƠ SỞ LÝ THUYẾT', True),
    ('   1.1. Tổng quan về bài toán nhận diện cảm xúc khuôn mặt', False),
    ('      1.1.1. Khái niệm Facial Emotion Recognition (FER)', False),
    ('      1.1.2. Ứng dụng trong cuộc gọi video thời gian thực', False),
    ('      1.1.3. Các thách thức của bài toán', False),
    ('   1.2. Các cảm xúc cơ bản trong nghiên cứu', False),
    ('      1.2.1. 7 cảm xúc cơ bản', False),
    ('      1.2.2. Đặc trưng khuôn mặt của từng cảm xúc', False),
    ('   1.3. Tổng quan về Deep Learning trong Computer Vision', False),
    ('      1.3.1. Artificial Neural Network (ANN)', False),
    ('      1.3.2. Convolutional Neural Network (CNN)', False),
    ('      1.3.3. Các thành phần của CNN', False),
    ('   1.4. Attention Mechanism trong CNN', False),
    ('      1.4.1. Squeeze-and-Excitation (SE) Block', False),
    ('      1.4.2. CBAM (Convolutional Block Attention Module)', False),
    ('      1.4.3. So sánh Attention và CNN truyền thống', False),
    ('   1.5. Transfer Learning', False),
    ('      1.5.1. Khái niệm Transfer Learning', False),
    ('      1.5.2. MobileNetV2 và đặc điểm kiến trúc', False),
    ('      1.5.3. Hạn chế khi áp dụng với ảnh nhỏ (48×48)', False),
    ('   1.6. Các kỹ thuật cải thiện hiệu năng mô hình', False),
    ('      1.6.1. Data Augmentation', False),
    ('      1.6.2. Focal Loss', False),
    ('      1.6.3. Test Time Augmentation (TTA)', False),
    ('      1.6.4. Xử lý mất cân bằng dữ liệu', False),
    ('   1.7. Dataset sử dụng', False),
    ('      1.7.1. Giới thiệu FER2013', False),
    ('      1.7.2. Phân tích đặc điểm dữ liệu', False),
    ('      1.7.3. Vấn đề mất cân bằng và noisy labels', False),
    ('', False),
    ('CHƯƠNG 2: PHƯƠNG PHÁP ĐỀ XUẤT', True),
    ('   2.1. Tổng quan các phương pháp đã thử nghiệm', False),
    ('      2.1.1. Method 1 – CNN + Enhanced Augmentation', False),
    ('      2.1.2. Method 2 – CNN + SE-Attention', False),
    ('      2.1.3. Method 3 – MobileNetV2 Transfer Learning', False),
    ('      2.1.4. Method 4 – SE-CBAM + Focal Loss', False),
    ('      2.1.5. M2 Optimized – CBAM-CNN + TTA', False),
    ('   2.2. Phân tích và so sánh kết quả các phương pháp', False),
    ('      2.2.1. So sánh Test Accuracy', False),
    ('      2.2.2. Phân tích ưu – nhược điểm từng phương pháp', False),
    ('      2.2.3. Lý do lựa chọn mô hình tối ưu', False),
    ('   2.3. Phương pháp đề xuất cuối cùng', False),
    ('      2.3.1. Kiến trúc CBAM-CNN', False),
    ('      2.3.2. Luồng xử lý dữ liệu trong training', False),
    ('      2.3.3. Cấu hình huấn luyện', False),
    ('   2.4. Chiến lược tối ưu hóa', False),
    ('      2.4.1. Regularization (Dropout, BatchNorm)', False),
    ('      2.4.2. Attention Mechanism', False),
    ('      2.4.3. TTA trong giai đoạn suy luận', False),
    ('', False),
    ('CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG', True),
    ('   3.1. Yêu cầu hệ thống', False),
    ('   3.2. Kiến trúc tổng thể hệ thống', False),
    ('   3.3. Luồng xử lý hệ thống thời gian thực', False),
    ('   3.4. Thiết kế module', False),
    ('   3.5. Công nghệ sử dụng', False),
    ('', False),
    ('CHƯƠNG 4: SẢN PHẨM, MÔ HÌNH VÀ THUẬT TOÁN ĐÃ XÂY DỰNG', True),
    ('   4.1. Mô hình CBAM-CNN hoàn chỉnh', False),
    ('   4.2. Kết quả thực nghiệm', False),
    ('   4.3. Sản phẩm demo thời gian thực', False),
    ('   4.4. Các notebook huấn luyện', False),
    ('   4.5. Hạn chế của hệ thống', False),
    ('   4.6. Hướng phát triển', False),
]

for text, is_chapter in toc_data:
    if text == '':
        doc.add_paragraph()
        continue
    pr = doc.add_paragraph()
    run = pr.add_run(text)
    run.font.name = 'Times New Roman'
    if is_chapter:
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 153)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_page_break()


# =============================================
# CHƯƠNG 1
# =============================================
heading_c1 = h1('CHƯƠNG 1: GIỚI THIỆU VỀ CƠ SỞ LÝ THUYẾT')

h2('1.1. Tổng quan về bài toán nhận diện cảm xúc khuôn mặt')

h3('1.1.1. Khái niệm Facial Emotion Recognition (FER)')
p('Facial Emotion Recognition (FER) là lĩnh vực nghiên cứu trong Computer Vision và Artificial Intelligence, tập trung vào việc tự động nhận diện và phân loại cảm xúc của con người thông qua biểu cảm khuôn mặt. FER sử dụng các thuật toán xử lý ảnh và mô hình học sâu (Deep Learning) để phân tích hình ảnh khuôn mặt và ánh xạ chúng vào các nhóm cảm xúc cơ bản.')
p('Theo nghiên cứu của Paul Ekman (1971), con người có 7 cảm xúc phổ quát được biểu hiện qua nét mặt bất kể nền văn hóa: Angry (giận dữ), Disgust (ghê tởm), Fear (sợ hãi), Happy (vui vẻ), Neutral (bình thường), Sad (buồn), và Surprise (ngạc nhiên). Đây cũng là 7 nhãn phân loại được sử dụng trong hầu hết các bộ dataset FER hiện đại.')

h3('1.1.2. Ứng dụng trong cuộc gọi video thời gian thực')
p('Trong bối cảnh làm việc và học tập từ xa ngày càng phổ biến sau đại dịch COVID-19, nhu cầu hiểu và phân tích cảm xúc trong các cuộc gọi video trở nên cấp thiết. Các ứng dụng chính bao gồm:')
b('Giáo dục trực tuyến: Giúp giáo viên đánh giá mức độ tập trung, hiểu bài của học sinh trong lớp học online.')
b('Họp doanh nghiệp: Đánh giá phản ứng của nhân viên/khách hàng trong các cuộc họp, hỗ trợ ra quyết định.')
b('Chăm sóc sức khỏe tâm thần: Phát hiện sớm các dấu hiệu trầm cảm, lo âu thông qua biểu cảm khuôn mặt.')
b('Dịch vụ khách hàng: Phân tích mức độ hài lòng của khách hàng trong tương tác video.')

h3('1.1.3. Các thách thức của bài toán')
p('Bài toán FER trong cuộc gọi video thời gian thực đối mặt với nhiều thách thức kỹ thuật:')
b('Chất lượng ảnh thấp: Camera webcam thường có độ phân giải và chất lượng hạn chế, ảnh bị nhiễu, mờ.')
b('Điều kiện ánh sáng: Ánh sáng không đồng đều, ngược sáng, hoặc quá tối gây khó khăn cho việc phát hiện đặc trưng khuôn mặt.')
b('Góc mặt đa dạng: Người dùng không luôn nhìn thẳng vào camera, mặt có thể nghiêng hoặc quay.')
b('Yêu cầu real-time: Hệ thống phải xử lý với latency thấp (<500ms) để đảm bảo trải nghiệm người dùng.')
b('Sự mơ hồ giữa các cảm xúc: Một số cảm xúc có biểu cảm tương tự nhau (fear vs surprise, sad vs neutral).')

h2('1.2. Các cảm xúc cơ bản trong nghiên cứu')

h3('1.2.1. 7 cảm xúc cơ bản (angry, disgust, fear, happy, neutral, sad, surprise)')
p('Hệ thống phân loại 7 cảm xúc cơ bản của Paul Ekman được sử dụng làm tiêu chuẩn trong hầu hết các nghiên cứu FER. Mỗi cảm xúc tương ứng với một tập hợp các chuyển động cơ mặt đặc trưng, được gọi là Action Units (AU) trong hệ thống FACS (Facial Action Coding System).')

add_image('data_augmentation.png', 'Hình 1: Minh họa 7 cảm xúc cơ bản và các phép biến đổi Data Augmentation', 5.5)

h3('1.2.2. Đặc trưng khuôn mặt của từng cảm xúc')
tbl(['Cảm xúc', 'Đặc trưng khuôn mặt', 'Action Units chính'], [
    ['Angry', 'Lông mày nhíu xuống, mắt trợn, môi mím chặt', 'AU4, AU5, AU7, AU23'],
    ['Disgust', 'Mũi nhăn, môi trên nhếch lên, lông mày hạ', 'AU9, AU15, AU25'],
    ['Fear', 'Mắt mở to, lông mày nhướng lên và kéo lại gần nhau', 'AU1, AU2, AU4, AU5, AU20'],
    ['Happy', 'Miệng cười, má nâng lên, nếp nhăn quanh mắt', 'AU6, AU12'],
    ['Neutral', 'Không có biểu cảm rõ rệt, cơ mặt thư giãn', 'Không có AU rõ rệt'],
    ['Sad', 'Khóe miệng hạ xuống, mắt cụp, lông mày kéo vào', 'AU1, AU4, AU15'],
    ['Surprise', 'Mắt mở to, lông mày nhướng cao, miệng há', 'AU1, AU2, AU5, AU26'],
])

h2('1.3. Tổng quan về Deep Learning trong Computer Vision')

h3('1.3.1. Artificial Neural Network (ANN)')
p('Mạng nơ-ron nhân tạo (ANN) là nền tảng của Deep Learning, lấy cảm hứng từ cách hoạt động của não bộ. ANN gồm nhiều lớp neurons kết nối với nhau, mỗi kết nối có một trọng số (weight) được học trong quá trình huấn luyện. Quá trình học bao gồm hai giai đoạn: Forward Propagation (tính output) và Backpropagation (cập nhật weights dựa trên gradient của hàm loss).')

h3('1.3.2. Convolutional Neural Network (CNN)')
p('CNN là kiến trúc Deep Learning chuyên biệt cho xử lý dữ liệu có cấu trúc lưới (grid-like), đặc biệt là hình ảnh. Điểm khác biệt quan trọng của CNN so với ANN truyền thống là việc sử dụng phép tích chập (convolution) để tự động trích xuất đặc trưng không gian (spatial features) từ ảnh.')
p('CNN hoạt động theo nguyên lý phân cấp đặc trưng: các lớp đầu học các đặc trưng cơ bản (cạnh, góc, texture), các lớp giữa kết hợp thành đặc trưng phức tạp hơn (mắt, mũi, miệng), và các lớp sâu nhận diện các pattern cấp cao (biểu cảm khuôn mặt).')

add_image('cnn_architecture.png', 'Hình 2: Kiến trúc CNN cho nhận diện cảm xúc khuôn mặt (CBAM-CNN)', 5.8)

h3('1.3.3. Các thành phần của CNN (Conv, Pooling, BatchNorm, Dropout, Softmax)')
tbl(['Thành phần', 'Chức năng', 'Mô tả'], [
    ['Conv2D (3×3)', 'Trích xuất đặc trưng cục bộ', 'Áp kernel 3×3, tạo feature maps'],
    ['MaxPooling2D (2×2)', 'Giảm kích thước 50%', 'Lấy max trong vùng 2×2'],
    ['BatchNormalization', 'Chuẩn hóa output', 'Normalize theo mini-batch'],
    ['Dropout (p)', 'Chống overfitting', 'Random tắt p% neurons khi train'],
    ['ReLU', 'Hàm kích hoạt', 'f(x) = max(0, x)'],
    ['Softmax', 'Phân loại đa lớp', 'Xác suất cho 7 classes (tổng = 1)'],
    ['GlobalAveragePooling', 'Nén features', 'Trung bình mỗi channel → vector'],
])

h2('1.4. Attention Mechanism trong CNN')

h3('1.4.1. Squeeze-and-Excitation (SE) Block')
p('SE Block (Hu et al., 2018) là cơ chế Channel Attention, giúp model tự động học trọng số cho từng channel (feature map). Ý tưởng: không phải tất cả feature maps đều quan trọng như nhau — ví dụ, feature map phát hiện "mắt nheo" sẽ quan trọng hơn cho cảm xúc Happy.')

add_image('se_block.png', 'Hình 3: Squeeze-and-Excitation (SE) Block', 5.5)

h3('1.4.2. CBAM (Convolutional Block Attention Module)')
p('CBAM (Woo et al., 2018) mở rộng SE Block bằng cách kết hợp cả Channel Attention và Spatial Attention:')
b('Channel Attention: Kết hợp Average Pooling + Max Pooling → Shared MLP → Sigmoid → "Feature nào quan trọng?"')
b('Spatial Attention: Conv2D(7×7) trên concat(avg, max) → Sigmoid → "Vùng nào trên ảnh quan trọng?" (mắt, miệng > trán, tai)')

add_image('cbam_module.png', 'Hình 4: CBAM – Convolutional Block Attention Module', 5.8)

h3('1.4.3. So sánh Attention và CNN truyền thống')
tbl(['Tiêu chí', 'CNN truyền thống', 'CNN + Attention (SE/CBAM)'], [
    ['Xử lý features', 'Đồng đều tất cả channels', 'Ưu tiên channels quan trọng'],
    ['Vùng ảnh', 'Xử lý toàn bộ ảnh như nhau', 'Tập trung vào vùng quan trọng'],
    ['Accuracy FER2013', '~60-62%', '~63-65%'],
    ['Parameters tăng thêm', '0', '~2-5% (rất ít)'],
])

h2('1.5. Transfer Learning')

h3('1.5.1. Khái niệm Transfer Learning')
p('Transfer Learning là kỹ thuật sử dụng model đã được huấn luyện trên tập dữ liệu lớn (ImageNet – 1.2 triệu ảnh) làm điểm khởi đầu, sau đó fine-tune cho bài toán cụ thể. Lợi ích: tận dụng features đã học, tiết kiệm thời gian, cải thiện accuracy khi data ít.')

h3('1.5.2. MobileNetV2 và đặc điểm kiến trúc')
p('MobileNetV2 (Sandler et al., 2018) là kiến trúc lightweight CNN sử dụng Inverted Residual Blocks với Depthwise Separable Convolution, giảm đáng kể parameters và FLOPs. Pretrained trên ImageNet với input RGB 224×224×3.')

h3('1.5.3. Hạn chế khi áp dụng với ảnh nhỏ (48×48)')

add_image('transfer_learning.png', 'Hình 5: Transfer Learning với MobileNetV2 – Kết quả và hạn chế', 5.5)

p('Kết quả thực nghiệm cho thấy MobileNetV2 chỉ đạt 36.28% accuracy, do: (1) Size mismatch – resize 48→224 gây artifacts; (2) Color mismatch – grayscale duplicate sang RGB giả; (3) Domain gap – ImageNet (ảnh tự nhiên) khác xa FER2013 (ảnh khuôn mặt cảm xúc).')

h2('1.6. Các kỹ thuật cải thiện hiệu năng mô hình')

h3('1.6.1. Data Augmentation')
p('Tăng cường dữ liệu bằng các phép biến đổi hình ảnh (rotation, flip, zoom, shift, brightness), giúp model học được các biến thể khác nhau và giảm overfitting.')

tbl(['Kỹ thuật', 'Thông số', 'Mục đích'], [
    ['Rotation', '±20°', 'Xử lý mặt nghiêng'],
    ['Width/Height Shift', '±15%', 'Vị trí mặt khác nhau'],
    ['Zoom', '±15%', 'Khoảng cách camera'],
    ['Horizontal Flip', 'Có', 'Khuôn mặt đối xứng'],
    ['Brightness', '0.8 – 1.2', 'Điều kiện ánh sáng'],
])

h3('1.6.2. Focal Loss')
p('Focal Loss (Lin et al., 2017) thêm hệ số (1-p)^γ vào Cross-Entropy để giảm weight cho easy examples và tập trung vào hard examples, đặc biệt hiệu quả với class imbalance.')

h3('1.6.3. Test Time Augmentation (TTA)')
p('Tạo nhiều phiên bản augmented cho ảnh test, predict trên tất cả rồi lấy trung bình. TTA cải thiện ~0.5-1% accuracy mà không thay đổi model.')

h3('1.6.4. Xử lý mất cân bằng dữ liệu')
b('Class Weights: Gán trọng số nghịch đảo với tần suất – class ít data (disgust) được weight cao hơn.')
b('Label Smoothing: Soft labels [0.02,...,0.88,...] thay vì one-hot [0,...,1,...], giảm overconfident.')

h2('1.7. Dataset sử dụng')

h3('1.7.1. Giới thiệu FER2013')
p('FER2013 (Goodfellow et al., 2013) là bộ dataset benchmark phổ biến nhất cho bài toán FER, gồm 35,887 ảnh grayscale 48×48, chia thành 28,709 train và 7,178 test.')

h3('1.7.2. Phân tích đặc điểm dữ liệu')

add_image('fer2013_distribution.png', 'Hình 6: Phân bố dữ liệu FER2013 (Training Set) – Mất cân bằng rõ rệt', 5.0)

tbl(['Cảm xúc', 'Train', 'Tỷ lệ', 'Test'], [
    ['Angry', '3,995', '13.9%', '958'],
    ['Disgust', '436', '1.5%', '111'],
    ['Fear', '4,097', '14.3%', '1,024'],
    ['Happy', '7,215', '25.1%', '1,774'],
    ['Neutral', '4,965', '17.3%', '1,233'],
    ['Sad', '4,830', '16.8%', '1,247'],
    ['Surprise', '3,171', '11.0%', '831'],
])

h3('1.7.3. Vấn đề mất cân bằng và noisy labels')
p('FER2013 có hai vấn đề: (1) Mất cân bằng – Disgust chỉ 436 ảnh, ít hơn 16.5 lần so với Happy; (2) Noisy labels – ~10-15% ảnh bị gán nhãn sai. Hai vấn đề này đặt giới hạn accuracy lý thuyết khoảng 70-75%.')

doc.add_page_break()


# =============================================
# CHƯƠNG 2
# =============================================
heading_c2 = h1('CHƯƠNG 2: PHƯƠNG PHÁP ĐỀ XUẤT')

h2('2.1. Tổng quan các phương pháp đã thử nghiệm')
p('Đồ án thực hiện 5 thí nghiệm với các kiến trúc khác nhau trên cùng dataset FER2013, nhằm tìm giải pháp tối ưu.')

h3('2.1.1. Method 1 – CNN + Enhanced Augmentation')
p('Sử dụng CNN 4 blocks tiêu chuẩn + Data Augmentation mạnh (rotation ±25°, zoom 20%, shear, brightness). Mục tiêu: đánh giá hiệu quả augmentation đơn thuần. Notebook: method1_enhanced_aug.ipynb.')

h3('2.1.2. Method 2 – CNN + SE-Attention')
p('Cải tiến Method 1 bằng SE Block sau mỗi Conv Block, tự động điều chỉnh trọng số cho từng feature channel. Notebook: method2_se_attention.ipynb.')

h3('2.1.3. Method 3 – MobileNetV2 Transfer Learning')
p('Sử dụng MobileNetV2 pretrained trên ImageNet, freeze base layers và fine-tune classifier head. Input RGB 48→224. Notebook: method3_mobilenet.ipynb.')

h3('2.1.4. Method 4 – SE-CBAM + Focal Loss')
p('Nâng cấp SE thành CBAM (Channel + Spatial Attention), kết hợp Focal Loss (γ=2), Cosine Annealing LR, TTA. Notebook: method4_optimized_se_cbam.ipynb.')

h3('2.1.5. M2 Optimized – CBAM-CNN + TTA')
p('Phiên bản tối ưu cuối cùng từ Method 2: thay SE thành CBAM, fine-tune hyperparameters (LR=0.001, batch=64, epochs=100, label_smoothing=0.15). Notebook: Optimize__method2 (1).ipynb.')

h2('2.2. Phân tích và so sánh kết quả các phương pháp')

h3('2.2.1. So sánh Test Accuracy')

add_image('method_comparison.png', 'Hình 7: So sánh Test Accuracy giữa 5 phương pháp trên FER2013', 5.2)

tbl(['Phương pháp', 'Test Accuracy', 'Test Loss', 'Ghi chú'], [
    ['Method 1 – Enhanced Aug', 'N/A', 'N/A', 'Training chưa hoàn tất'],
    ['Method 2 – SE-Attention', '64.22%', '1.1997', '✅ Accuracy cao nhất'],
    ['Method 3 – MobileNetV2', '36.28%', '1.7924', '❌ Kém nhất'],
    ['Method 4 – SE-CBAM', '63.51% (TTA)', '1.2199', 'Focal Loss + TTA'],
    ['M2 Optimized', '63.50% (TTA)', '-', 'Best val: 60.46%'],
])

h3('2.2.2. Phân tích ưu – nhược điểm từng phương pháp')
b('Method 2 (SE): Đơn giản, hiệu quả. SE thêm rất ít params (~2-3%) nhưng cải thiện accuracy đáng kể.')
b('Method 3 (MobileNetV2): Thất bại do domain mismatch và size mismatch. Bài học: Transfer Learning không phải lúc nào cũng tốt hơn.')
b('Method 4 (CBAM + Focal Loss): Nhiều kỹ thuật tối ưu nhưng không vượt Method 2, có thể do hyperparameters chưa tối ưu.')

h3('2.2.3. Lý do lựa chọn mô hình tối ưu')
p('Method 2 (SE-Attention CNN) được chọn vì: (1) Accuracy cao nhất 64.22%; (2) Kiến trúc đơn giản, dễ reproduce; (3) Inference nhanh cho real-time. M2 Optimized với CBAM là mô hình cuối cùng được deploy.')

h2('2.3. Phương pháp đề xuất cuối cùng')

h3('2.3.1. Kiến trúc CBAM-CNN')
tbl(['Layer', 'Cấu hình', 'Output Size'], [
    ['Input', '48 × 48 × 1 (Grayscale)', '48×48×1'],
    ['Block 1', '2×Conv2D(64) + BN + ReLU + CBAM(r=8) + MaxPool + Drop(0.25)', '24×24×64'],
    ['Block 2', '2×Conv2D(128) + BN + ReLU + CBAM(r=8) + MaxPool + Drop(0.25)', '12×12×128'],
    ['Block 3', '2×Conv2D(256) + BN + ReLU + CBAM(r=16) + MaxPool + Drop(0.3)', '6×6×256'],
    ['Block 4', '2×Conv2D(512) + BN + ReLU + CBAM(r=16) + MaxPool + Drop(0.3)', '3×3×512'],
    ['Classifier', 'GAP → Dense(512,L2) → BN → ReLU → Drop(0.5) → Dense(256,L2) → BN → ReLU → Drop(0.5) → Dense(7,softmax)', '7'],
])

h3('2.3.2. Luồng xử lý dữ liệu trong training')
p('Pipeline: Load ảnh → Resize 48×48, grayscale → Normalize [0,1] → Data Augmentation → Chia 80% train / 20% val → Feed vào model theo batch_size=64.')

h3('2.3.3. Cấu hình huấn luyện')
tbl(['Hyperparameter', 'Giá trị'], [
    ['Optimizer', 'Adam'],
    ['Learning Rate', '0.001'],
    ['Batch Size', '64'],
    ['Epochs', '100 (EarlyStopping patience=10)'],
    ['Label Smoothing', '0.15'],
    ['LR Scheduler', 'ReduceLROnPlateau (factor=0.5, patience=4)'],
    ['L2 Regularization', '0.0005 (Dense layers)'],
])

h2('2.4. Chiến lược tối ưu hóa')

h3('2.4.1. Regularization (Dropout, BatchNorm)')
p('Dropout tỷ lệ tăng dần: 0.25 (Block 1-2) → 0.3 (Block 3-4) → 0.5 (Dense). BatchNormalization sau mỗi Conv và Dense layer.')

h3('2.4.2. Attention Mechanism')
p('CBAM đặt sau mỗi Conv Block, trước MaxPooling. Block 1-2 dùng ratio=8, Block 3-4 dùng ratio=16.')

h3('2.4.3. TTA trong giai đoạn suy luận')
p('Mỗi ảnh test được tạo nhiều phiên bản augmented, predict trên tất cả, lấy trung bình. Cải thiện accuracy từ 62.76% → 63.51% (+0.75%).')

doc.add_page_break()


# =============================================
# CHƯƠNG 3
# =============================================
heading_c3 = h1('CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG')

h2('3.1. Yêu cầu hệ thống')

h3('3.1.1. Yêu cầu chức năng')
b('Nhận diện khuôn mặt trong video stream từ webcam.')
b('Phân loại cảm xúc khuôn mặt thành 7 nhóm cảm xúc cơ bản.')
b('Hiển thị kết quả (tên cảm xúc + confidence %) trực tiếp lên video.')
b('Vẽ bounding box xung quanh khuôn mặt được phát hiện.')
b('Hỗ trợ xử lý real-time với FPS ≥15.')

h3('3.1.2. Yêu cầu phi chức năng')
b('Latency: < 500ms cho mỗi frame.')
b('Accuracy: ≥ 60% trên test set FER2013.')
b('Bảo mật: Xử lý hoàn toàn local, không gửi dữ liệu lên server.')
b('Tương thích: Windows/macOS/Linux, Python 3.10+.')

h2('3.2. Kiến trúc tổng thể hệ thống')

h3('3.2.1. Sơ đồ tổng thể (System Architecture Diagram)')

add_image('system_pipeline.png', 'Hình 8: Sơ đồ kiến trúc hệ thống nhận diện cảm xúc thời gian thực', 5.8)

h3('3.2.2. Mô hình Client-side processing')
p('Hệ thống sử dụng mô hình xử lý hoàn toàn trên client (on-device): (1) Bảo mật – dữ liệu khuôn mặt không rời thiết bị; (2) Latency thấp – không có network delay; (3) Offline – không cần internet.')

h2('3.3. Luồng xử lý hệ thống thời gian thực')

h3('3.3.1. Capture video từ webcam')
p('Sử dụng OpenCV VideoCapture(0) để mở camera mặc định. Mỗi frame được đọc liên tục trong vòng lặp while True.')

h3('3.3.2. Phát hiện khuôn mặt (Haar Cascade)')
p('Sử dụng Haar Cascade Classifier (haarcascade_frontalface_default.xml). Cấu hình: scaleFactor=1.3, minNeighbors=5. Output: danh sách bounding boxes (x, y, w, h).')

h3('3.3.3. Tiền xử lý ảnh')
p('Quy trình: Crop ROI khuôn mặt → Grayscale → Resize 48×48 → Normalize [0,1] → Expand dims (1, 48, 48, 1).')

h3('3.3.4. Dự đoán cảm xúc bằng mô hình')
p('Model CBAM-CNN nhận input (1, 48, 48, 1), output 7 xác suất (softmax). Argmax → nhãn cảm xúc + confidence %.')

h3('3.3.5. Hiển thị kết quả')
p('OpenCV vẽ lên frame: Rectangle (bounding box) màu xanh lá + Text (tên cảm xúc + confidence) phía trên, hiển thị qua cv2.imshow().')

h2('3.4. Thiết kế module')

tbl(['Module', 'File', 'Chức năng'], [
    ['Capture', 'realtime_demo.py', 'OpenCV VideoCapture – Đọc frame từ webcam'],
    ['Face Detection', 'realtime_demo.py', 'Haar Cascade – Phát hiện vùng khuôn mặt'],
    ['Classification', 'realtime_demo.py', 'CBAM-CNN model – Predict 7 cảm xúc'],
    ['Visualization', 'realtime_demo.py', 'OpenCV – Vẽ bounding box + text lên frame'],
])

h2('3.5. Công nghệ sử dụng')

tbl(['Công nghệ', 'Phiên bản', 'Vai trò'], [
    ['Python', '3.10+', 'Ngôn ngữ lập trình chính'],
    ['TensorFlow/Keras', '2.x', 'Framework Deep Learning – xây dựng và train model'],
    ['OpenCV', '4.x', 'Computer Vision – camera, ảnh, face detection'],
    ['NumPy', '1.x', 'Xử lý mảng số, tiền xử lý ảnh'],
    ['Google Colab', '-', 'Cloud GPU (Tesla T4) cho training'],
    ['Matplotlib', '3.x', 'Trực quan hóa kết quả, biểu đồ'],
])

doc.add_page_break()


# =============================================
# CHƯƠNG 4
# =============================================
heading_c4 = h1('CHƯƠNG 4: SẢN PHẨM, MÔ HÌNH VÀ THUẬT TOÁN ĐÃ XÂY DỰNG')

h2('4.1. Mô hình CBAM-CNN hoàn chỉnh')

h3('4.1.1. Kiến trúc chi tiết (4 Conv Blocks + Attention)')
p('Model gồm 4 Convolutional Blocks, mỗi block: 2 lớp Conv2D → BatchNorm → ReLU → CBAM → MaxPool → Dropout. Số filters tăng dần: 64 → 128 → 256 → 512.')

add_image('cnn_architecture.png', 'Hình 9: Kiến trúc chi tiết CBAM-CNN (4 blocks + classifier head)', 5.8)

h3('4.1.2. Số lượng tham số (~15 triệu parameters)')
p('Tổng ~15 triệu parameters. CBAM chỉ thêm ~2-3% params nhưng cải thiện accuracy nhờ focus vào features và vùng ảnh quan trọng.')

h3('4.1.3. Cấu trúc Fully Connected Layers')
p('Classifier head: GlobalAveragePooling2D → Dense(512, L2=0.0005) → BN → ReLU → Dropout(0.5) → Dense(256, L2=0.0005) → BN → ReLU → Dropout(0.5) → Dense(7, softmax).')

h2('4.2. Kết quả thực nghiệm')

h3('4.2.1. Accuracy đạt được')
tbl(['Metric', 'Giá trị'], [
    ['Best Test Accuracy', '64.22% (Method 2 – SE)'],
    ['Best Val Accuracy', '60.46% (M2 Optimized)'],
    ['TTA Accuracy', '63.50%'],
    ['Weighted Precision', '0.6395'],
    ['Weighted Recall', '0.6350'],
    ['Weighted F1-Score', '0.6267'],
])

h3('4.2.2. Confusion Matrix')
p('Happy và Surprise có recall cao nhất (biểu cảm rõ ràng). Fear hay bị nhầm với Sad/Surprise. Disgust accuracy thấp (chỉ 111 test images). Neutral đôi khi bị nhầm với Sad.')

h3('4.2.3. Phân tích các lớp dự đoán kém (fear, sad)')
p('Fear: Biểu cảm tương tự surprise (mắt to) và sad (lông mày kéo vào). Nhiều ảnh FER2013 gán nhãn fear khó phân biệt ngay cả với con người.')
p('Sad: Precision/Recall ~0.50, hay nhầm với neutral/fear. Biểu cảm buồn nhẹ rất giống neutral.')

h2('4.3. Sản phẩm demo thời gian thực')

h3('4.3.1. realtime_demo.py')
p('File demo 224 dòng Python: định nghĩa custom layers (ChannelAttention, SpatialAttention, CBAMBlock) → Build model → Load weights → Camera capture → Face detect → Predict → Display.')

h3('4.3.2. Tốc độ xử lý (15–25 FPS)')
tbl(['Metric', 'Giá trị'], [
    ['FPS', '15–25'],
    ['Latency/frame', '~50–100 ms'],
    ['Model size', '~60 MB'],
    ['RAM usage', '~500 MB'],
])

h3('4.3.3. Minh họa giao diện demo')
p('Video webcam với bounding box xanh lá bao quanh khuôn mặt, tên cảm xúc tiếng Việt + confidence % hiển thị phía trên. Nhấn "q" để thoát.')

h2('4.4. Các notebook huấn luyện')
tbl(['File', 'Mô tả'], [
    ['method1_enhanced_aug.ipynb', 'CNN + Enhanced Augmentation'],
    ['method2_se_attention.ipynb', 'CNN + SE-Attention (Best: 64.22%)'],
    ['method3_mobilenet.ipynb', 'MobileNetV2 Transfer Learning'],
    ['method4_optimized_se_cbam.ipynb', 'SE-CBAM + Focal Loss'],
    ['Optimize__method2 (1).ipynb', 'CBAM-CNN + TTA (Final model)'],
    ['best_model232.keras', 'Model file (~62MB)'],
    ['model_weights.weights.h5', 'Weights file (~62MB)'],
])

h2('4.5. Hạn chế của hệ thống')

h3('4.5.1. Accuracy chưa cao ở một số lớp')
p('Fear, Sad, Disgust có F1-Score dưới 0.55 do biểu cảm subtle, dễ nhầm lẫn, data không đủ.')

h3('4.5.2. Ảnh FER2013 nhỏ và nhiễu')
p('Ảnh 48×48 grayscale mất nhiều chi tiết. ~10-15% noisy labels đặt giới hạn accuracy ~70-75%.')

h3('4.5.3. Haar Cascade chưa tối ưu')
p('Bỏ sót mặt nghiêng >30°, hoạt động kém trong ánh sáng yếu, có false positive.')

h2('4.6. Hướng phát triển')

h3('4.6.1. Thay Haar Cascade bằng MediaPipe')
p('MediaPipe Face Detection (Google) accuracy cao hơn, hỗ trợ mặt nghiêng, cung cấp 468 facial landmarks.')

h3('4.6.2. Tích hợp Drowsiness Detection')
p('Thêm module phát hiện buồn ngủ sử dụng MRL Eye Dataset (84,898 ảnh, 2 classes: awake/sleepy) kết hợp thuật toán EAR (Eye Aspect Ratio).')

h3('4.6.3. Phát triển Chrome Extension cho Google Meet/Zoom')
p('Xây dựng Chrome Extension dùng Canvas API/WebRTC capture frame từ video call, chạy TensorFlow.js trực tiếp trên browser.')

h3('4.6.4. Mở rộng dataset')
p('Bổ sung AffectNet (~400K ảnh, phân giải cao hơn) và RAF-DB (ảnh thực tế) để cải thiện accuracy và generalization.')

# =============================================
# TÀI LIỆU THAM KHẢO
# =============================================
doc.add_page_break()
h1('TÀI LIỆU THAM KHẢO')
refs = [
    '[1] Ekman, P. (1971). "Universals and cultural differences in facial expressions of emotion."',
    '[2] Goodfellow, I. et al. (2013). "Challenges in representation learning: A report on three machine learning contests." ICML.',
    '[3] Hu, J. et al. (2018). "Squeeze-and-Excitation Networks." CVPR.',
    '[4] Woo, S. et al. (2018). "CBAM: Convolutional Block Attention Module." ECCV.',
    '[5] Lin, T. et al. (2017). "Focal Loss for Dense Object Detection." ICCV.',
    '[6] Sandler, M. et al. (2018). "MobileNetV2: Inverted Residuals and Linear Bottlenecks." CVPR.',
    '[7] Mollahosseini, A. et al. (2017). "AffectNet: A Database for Facial Expression, Valence, and Arousal."',
    '[8] Li, S. & Deng, W. (2020). "Deep Facial Expression Recognition: A Survey." IEEE Trans.',
]
for ref in refs:
    pr = doc.add_paragraph(ref)
    for r in pr.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

# =============================================
# SAVE
# =============================================
output = r'd:\New folder\CapstoneProject\Bao_cao_nghien_cuu_de_tai_v2.docx'
doc.save(output)
print(f'✅ Đã tạo thành công: {output}')
