# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

def para(text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
            for r in t.rows[ri+1].cells[ci].paragraphs[0].runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    doc.add_paragraph()

# === TRANG BÌA ===
for _ in range(4):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TÓM TẮT ĐỒ ÁN TỐT NGHIỆP\n\nXÂY DỰNG HỆ THỐNG NHẬN DIỆN CẢM XÚC KHUÔN MẶT\nTHEO THỜI GIAN THỰC TRONG CUỘC GỌI VIDEO')
run.bold = True; run.font.size = Pt(18); run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)
doc.add_page_break()

# === 1. MỤC TIÊU ===
heading('1. Mục tiêu đề tài')
para('Xây dựng hệ thống nhận diện 7 cảm xúc cơ bản (angry, disgust, fear, happy, neutral, sad, surprise) từ khuôn mặt theo thời gian thực, ứng dụng trong cuộc gọi video trực tuyến.')

# === 2. DATASET ===
heading('2. Dataset')
para('FER2013: 35,887 ảnh grayscale 48×48px, chia thành 28,709 ảnh train và 7,178 ảnh test, 7 lớp cảm xúc. Dữ liệu mất cân bằng (disgust chỉ ~1.5%).')

# === 3. CÁC PHƯƠNG PHÁP ===
heading('3. Các phương pháp đã thử nghiệm')
table(
    ['Method', 'Kỹ thuật chính', 'Test Accuracy'],
    [
        ['1 - Enhanced Aug', 'CNN + Augmentation mạnh', 'Chưa hoàn tất'],
        ['2 - SE-Attention', 'CNN + Squeeze-Excitation', '64.22%'],
        ['3 - MobileNetV2', 'Transfer Learning (RGB)', '36.28%'],
        ['4 - SE-CBAM', 'CNN + CBAM + Focal Loss + TTA', '63.51%'],
        ['M2 Optimized', 'CBAM-CNN + TTA (Final)', '63.50%'],
    ]
)
para('→ Method 2 (SE-Attention) đạt kết quả tốt nhất: 64.22%. MobileNetV2 kém do ảnh FER2013 quá nhỏ (48×48) cho model pretrained.')

# === 4. KIẾN TRÚC MODEL TỐI ƯU ===
heading('4. Kiến trúc model (CBAM-CNN)')
para('4 blocks Conv2D (64→128→256→512) + BatchNorm + ReLU + CBAM Attention + MaxPool + Dropout, kết thúc bằng GlobalAveragePooling → Dense(512) → Dense(256) → Dense(7, softmax). Tổng ~15 triệu params.')

# === 5. SẢN PHẨM ===
heading('5. Sản phẩm đã xây dựng')
bullet('realtime_demo.py: Demo nhận diện cảm xúc qua webcam (15-25 FPS)')
bullet('Model CBAM-CNN đã train (~60MB weights)')
bullet('5 notebooks Google Colab (4 methods + 1 optimized)')
bullet('Hệ thống: OpenCV capture → Haar Cascade detect face → Model predict → Hiển thị kết quả')

# === 6. HẠN CHẾ & HƯỚNG PHÁT TRIỂN ===
heading('6. Hạn chế & Hướng phát triển')
bullet('Hạn chế: Accuracy một số cảm xúc khó (fear, sad) còn thấp; FER2013 có ~10-15% noisy labels')
bullet('Phát triển: Thay Haar Cascade bằng MediaPipe; thêm Drowsiness Detection (MRL Eye Dataset); tích hợp Chrome Extension cho Google Meet/Zoom')

# === SAVE ===
path = r'd:\New folder\CapstoneProject\Tom_tat_do_an.docx'
doc.save(path)
print(f'✅ Đã tạo: {path}')
